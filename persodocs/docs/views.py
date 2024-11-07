from django.shortcuts import render, redirect, get_object_or_404
from .forms import DocForm, UploadExcelForm, UserRegistrationForm
from .models import Doc
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse, HttpResponseBadRequest
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from pathlib import Path
import zipfile
import io
import pandas as pd
import re
import os
from django.conf import settings
import uuid
import glob
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm

# Create your views here.

#display the uploaded templates names & word doc upload form
@login_required
def home(request):
    if request.method == 'POST':
        form = DocForm(request.POST, request.FILES)
        if form.is_valid():
            new_file = form.save(commit=False)
            new_file.user = request.user
            new_file.save()
            return redirect('home')
    else:
        form = DocForm()
        
    docs = Doc.objects.all().order_by('-uploaded_at')
    return render(request, 'home.html', {'DocForm': form, 'docs': docs})


#display the uploaded document contents & excel upload form
@login_required
def doc_detail(request, doc_id):
    
    #render the excel upload form
    form = UploadExcelForm()
    
    
    doc = get_object_or_404(Doc, pk=doc_id)
    
    #parse the doc file content
    doc_path = doc.file.path
    doc_content = []
    if doc_path.endswith('.docx'):
        document = docx.Document(doc_path)
        for paragraph in document.paragraphs:
            # Initialize the paragraph data
            paragraph_content = []
            
            # Handle alignment (left, center, right, justified)
            align = paragraph.alignment
            align_class = ""
            if align == WD_ALIGN_PARAGRAPH.CENTER:
                align_class = "text-center"
            elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                align_class = "text-right"
            elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                align_class = "text-justify"
            # You can apply "text-left" as default (for left-aligned text)

            for run in paragraph.runs:
                run_text = run.text
                
                # Check for bold and apply <strong> tag
                if run.bold:
                    run_text = f"<strong>{run_text}</strong>"
                
                # Check for italic and apply <em> tag
                if run.italic:
                    run_text = f"<em>{run_text}</em>"
                
                # Check for underline and apply <u> tag
                if run.underline:
                    run_text = f"<u>{run_text}</u>"
                
                # Handle font size (in points) and apply to style
                font_size = None
                if run.font.size:
                    font_size = run.font.size.pt  # Get font size in points
                    run_text = f"<span style='font-size:{font_size}px'>{run_text}</span>"
                
                # Handle font color (if it exists)
                color = None
                if run.font.color and run.font.color.rgb:
                    color = run.font.color.rgb
                    color_hex = f"#{color[0]:02x}{color[1]:02x}{color[2]:02x}"  # Convert RGB to hex
                    run_text = f"<span style='color:{color_hex}'>{run_text}</span>"
                
                # Handle font family (if it exists)
                font_family = run.font.name if run.font.name else 'Arial'  # Default to Arial if no font is set
                run_text = f"<span style='font-family:{font_family};'>{run_text}</span>"
                
                # Add the formatted text to the paragraph content
                paragraph_content.append(run_text)
            
            # Combine all the runs into one paragraph and apply alignment class
            doc_content.append(f"<p class='{align_class}'>" + "".join(paragraph_content) + "</p>")    
    return render(request, 'doc_detail.html', {'doc': doc, 'doc_content': doc_content, 'form': form})


#read the excel and generate the personalized documents
@login_required
def generate_doc(request, doc_id):
    if request.method == 'POST':
        print("Inside doc generation")
        doc = get_object_or_404(Doc, pk=doc_id)
        
        # Handle the uploaded Excel file
        form = UploadExcelForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES['excel_file']
            
            # Generate a unique filename for the temp Excel file
            temp_filename = f'temp_{uuid.uuid4()}.xlsx'
            temp_path = default_storage.save(temp_filename, ContentFile(excel_file.read()))
            
            # Load data from Excel file
            customer_data = pd.read_excel(default_storage.path(temp_path))
            
            # Load the specific Word template associated with this doc_id
            template_path = doc.file.path
            template = docx.Document(template_path)
            
            # Create an in-memory zip file to store generated documents
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Generate personalized documents for each customer
                for _, row in customer_data.iterrows():
                    doc_copy = docx.Document(template_path)
                    
                    # Process each paragraph to replace placeholders
                    for paragraph in doc_copy.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            
                            # Find all placeholders within {{ }}
                            placeholders = re.findall(r"\{\{(.*?)\}\}", run_text)
                            print(f"Found placeholders in paragraph: {placeholders}")

                            for placeholder in placeholders:
                                # Get the value from the row if it exists, default to empty if not found
                                value = str(row.get(placeholder.strip(), ""))
                                print(f"Replacing placeholder '{{{{ {placeholder} }}}}' with value '{value}'")

                                # Replace the placeholder in the run text
                                run_text = run_text.replace(f"{{{{{placeholder}}}}}", value)
                            
                            # Update the run text with the replaced values
                            run.text = run_text

                    # Save the generated document in memory
                    output_file_name = f"{row['name']}_personalized.docx"
                    temp_output = io.BytesIO()
                    doc_copy.save(temp_output)
                    temp_output.seek(0)
                    zip_file.writestr(output_file_name, temp_output.read())

            # Return the zip file as a downloadable response
            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="personalized_documents.zip"'
            return response

    # If GET or form invalid, redirect back to `doc_detail`
    print("Docs generated view")
    return redirect('doc_detail', doc_id=doc_id)

#update the uploaded templates
@login_required
def update_doc(request, doc_id):
    doc = get_object_or_404(Doc, pk=doc_id)
    if request.method == 'POST':
        form = DocForm(request.POST, request.FILES, instance=doc)
        if form.is_valid():
            form.save()
            return redirect('home')
    else:
        form = DocForm(instance=doc)
    return render(request, 'doc_update.html', {'DocForm': form, 'doc': doc})

#delete the uploaded templates
login_required
def delete_doc(request, doc_id):
    doc = get_object_or_404(Doc, pk=doc_id)
    if request.method == 'POST':
        # Delete the associated document file if it exists
        if doc.file and os.path.exists(doc.file.path):
            os.remove(doc.file.path)
        
        # Delete the temporary Excel file if it exists
        temp_dir = settings.MEDIA_ROOT  # or another directory where temporary files are stored
        temp_files_pattern = os.path.join(temp_dir, 'temp_*.xlsx')
        # Find and delete all matching temp files
        for temp_file in glob.glob(temp_files_pattern):
            try:
                os.remove(temp_file)
            except OSError as e:
                print(f"Error deleting file {temp_file}: {e}")
        
        # Delete the Doc instance from the database
        doc.delete()
        return redirect('home')
    return render(request, 'doc_delete.html', {'doc': doc})


#user guide view
def user_guide(request):
    return render(request, 'user_guide.html')
    # pass
    

#authentication and authorization
def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password1'])
            user.save()
            login(request, user)
            return redirect('home')
    else:
        form = UserRegistrationForm()
    return render(request, 'registration/registration.html', {'form': form})
