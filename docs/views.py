from django.shortcuts import render, redirect, get_object_or_404
from .forms import DocForm, UploadExcelForm, UserRegistrationForm
from .models import Doc
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse, HttpResponseBadRequest
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from pathlib import Path
import zipfile
import io
import pandas as pd
import re
import os
from django.conf import settings
import uuid
import glob
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm
import boto3


# Get the S3 client
s3_client = boto3.client('s3', 
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
                region_name=settings.AWS_S3_REGION_NAME)

# Create your views here.

#display the uploaded templates names & word doc upload form
@login_required
def home(request):
    if request.method == 'POST':
        form = DocForm(request.POST, request.FILES)
        if form.is_valid():
            new_file = form.save(commit=False)
            new_file.user = request.user
            new_file.save()
            return redirect('home')
    else:
        form = DocForm()
        
    docs = Doc.objects.all().order_by('-uploaded_at')
    print(docs)
    return render(request, 'home.html', {'DocForm': form, 'docs': docs})


#display the uploaded document contents & excel upload form
@login_required
def doc_detail(request, doc_id):
    # Render the excel upload form
    form = UploadExcelForm()

    # Get the document object
    doc = get_object_or_404(Doc, pk=doc_id)

    # Get the S3 URL of the document
    s3_url = doc.file.url  # This will give you the S3 URL of the file

    # Download the file from S3
    file_obj = s3_client.get_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=doc.file.name)

    # Read the content of the file into memory
    file_content = file_obj['Body'].read()

    # Process the content if it's a DOCX file
    doc_content = []
    if s3_url.endswith('.docx'):
        # Use io.BytesIO to treat the file content as a file-like object
        docx_file = io.BytesIO(file_content)
        document = docx.Document(docx_file)
        
        for paragraph in document.paragraphs:
            paragraph_content = []
            
            align_class = ""
            align = paragraph.alignment
            if align == WD_ALIGN_PARAGRAPH.CENTER:
                align_class = "text-center"
            elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                align_class = "text-right"
            elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                align_class = "text-justify"

            for run in paragraph.runs:
                run_text = run.text

                # Check for bold, italic, underline, and font properties
                if run.bold:
                    run_text = f"<strong>{run_text}</strong>"
                if run.italic:
                    run_text = f"<em>{run_text}</em>"
                if run.underline:
                    run_text = f"<u>{run_text}</u>"
                
                font_size = None
                if run.font.size:
                    font_size = run.font.size.pt
                    run_text = f"<span style='font-size:{font_size}px'>{run_text}</span>"
                
                color = None
                if run.font.color and run.font.color.rgb:
                    color = run.font.color.rgb
                    color_hex = f"#{color[0]:02x}{color[1]:02x}{color[2]:02x}"
                    run_text = f"<span style='color:{color_hex}'>{run_text}</span>"
                
                font_family = run.font.name if run.font.name else 'Arial'
                run_text = f"<span style='font-family:{font_family};'>{run_text}</span>"
                
                paragraph_content.append(run_text)

            doc_content.append(f"<p class='{align_class}'>" + "".join(paragraph_content) + "</p>")

    return render(request, 'doc_detail.html', {'doc': doc, 'doc_content': doc_content, 'form': form})


#read the excel and generate the personalized documents
@login_required
def generate_doc(request, doc_id):
    if request.method == 'POST':
        print("Inside doc generation")
        doc = get_object_or_404(Doc, pk=doc_id)
        
        # Handle the uploaded Excel file
        form = UploadExcelForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES['excel_file']
            
            # Load data from the uploaded Excel file directly in memory
            customer_data = pd.read_excel(excel_file)
            
            # Open the specific Word template file associated with this doc_id directly in memory
            with default_storage.open(doc.file.name, 'rb') as template_file:
                template_content = template_file.read()

            # Load the Word template as an in-memory document
            template = docx.Document(io.BytesIO(template_content))
            
            # Create an in-memory zip file to store generated documents
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Generate personalized documents for each customer
                for _, row in customer_data.iterrows():
                    # Create a new document from the template for each row in Excel
                    doc_copy = docx.Document(io.BytesIO(template_content))
                    
                    # Process each paragraph to replace placeholders
                    for paragraph in doc_copy.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            
                            # Find all placeholders within {{ }}
                            placeholders = re.findall(r"\{\{(.*?)\}\}", run_text)
                            print(f"Found placeholders in paragraph: {placeholders}")

                            for placeholder in placeholders:
                                # Get the value from the row if it exists, default to empty if not found
                                value = str(row.get(placeholder.strip(), ""))
                                print(f"Replacing placeholder '{{{{ {placeholder} }}}}' with value '{value}'")

                                # Replace the placeholder in the run text
                                run_text = run_text.replace(f"{{{{{placeholder}}}}}", value)
                            
                            # Update the run text with the replaced values
                            run.text = run_text

                    # Save the generated document in memory
                    output_file_name = f"{row['name']}_personalized.docx"
                    temp_output = io.BytesIO()
                    doc_copy.save(temp_output)
                    temp_output.seek(0)
                    zip_file.writestr(output_file_name, temp_output.read())

            # Return the zip file as a downloadable response
            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="personalized_documents.zip"'
            return response

    # If GET or form invalid, redirect back to `doc_detail`
    print("Docs generated view")
    return redirect('doc_detail', doc_id=doc_id)

#update the uploaded templates
@login_required
def update_doc(request, doc_id):
    # Retrieve the document instance from the database
    doc = get_object_or_404(Doc, pk=doc_id)

    # Process the form submission
    if request.method == 'POST':
        form = DocForm(request.POST, request.FILES, instance=doc)
        if form.is_valid():
            try:
                # Save the form, and the file will be saved to S3 automatically
                form.save()
                return redirect('home')
            except Exception as e:
                print(f"Error saving document: {e}")
                return render(request, 'doc_update.html', {
                    'DocForm': form,
                    'doc': doc,
                    'error_message': f"An error occurred while saving the document: {e}"
                })
    else:
        # If not a POST request, just load the form with the current doc instance
        form = DocForm(instance=doc)
    return render(request, 'doc_update.html', {'DocForm': form, 'doc': doc})

#delete the uploaded templates
login_required
def delete_doc(request, doc_id):
    # Retrieve the document instance from the database
    doc = get_object_or_404(Doc, pk=doc_id)
    
    if request.method == 'POST':
        # Delete the associated document file on S3 if it exists
        if doc.file:
            doc.file.delete(save=False)  # Deletes the file on S3 without saving the model again

        # Delete any temporary Excel files if they exist
        temp_dir = settings.MEDIA_ROOT  # Assuming temporary files are stored here
        temp_files_pattern = os.path.join(temp_dir, 'temp_*.xlsx')
        for temp_file in glob.glob(temp_files_pattern):
            try:
                os.remove(temp_file)
            except OSError as e:
                print(f"Error deleting temporary file {temp_file}: {e}")

        # Delete the Doc instance from the database
        doc.delete()
        return redirect('home')

    return render(request, 'doc_delete.html', {'doc': doc})

#user guide view
def user_guide(request):
    return render(request, 'user_guide.html')
    # pass
    

#authentication and authorization
def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password1'])
            user.save()
            login(request, user)
            return redirect('home')
    else:
        form = UserRegistrationForm()
    return render(request, 'registration/registration.html', {'form': form})
